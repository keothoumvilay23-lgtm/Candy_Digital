# -*- coding: utf-8 -*-
"""Generate updated graduation thesis report for Candy Digital."""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import sys

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
sys.path.insert(0, SCRIPT_DIR)

OUTPUT = r'c:\Users\Admin\Downloads\BaoCao_CandyDigital_v4_HoanThien.docx'
DIAGRAM_DIR = os.path.join(SCRIPT_DIR, 'diagrams')

def set_cell_shading(cell, color="D9E2F3"):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.bold = bold
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i])
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
    doc.add_paragraph()
    return table

def center_para(doc, text, size=14, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_image(doc, image_path, caption, width=Inches(6.2)):
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=width)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.italic = True
    else:
        add_para(doc, f'[{caption} — chưa tìm thấy file ảnh]', indent=True)

def add_code_block(doc, title, code, filename=''):
    add_para(doc, title, bold=True)
    if filename:
        add_para(doc, f'File: {filename}', indent=False)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    doc.add_paragraph()

def read_source(relative_path, start=1, end=None):
    path = os.path.join(PROJECT_ROOT, relative_path.replace('/', os.sep))
    if not os.path.exists(path):
        return f'// Khong tim thay file: {relative_path}'
    with open(path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    if end is None:
        end = len(lines)
    selected = lines[start - 1:end]
    numbered = []
    for i, line in enumerate(selected, start=start):
        numbered.append(f'{i:4d}| {line.rstrip()}')
    return '\n'.join(numbered)

def build_document():
    from generate_diagrams import generate_all
    generate_all()
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)

    # ===== TRANG BÌA =====
    for _ in range(3):
        doc.add_paragraph()
    center_para(doc, 'TRƯỜNG ĐẠI HỌC [TÊN TRƯỜNG]', 14, True)
    center_para(doc, 'KHOA CÔNG NGHỆ THÔNG TIN', 14, True)
    center_para(doc, 'Bộ môn: Công nghệ phần mềm', 13)
    for _ in range(4):
        doc.add_paragraph()
    center_para(doc, 'BÁO CÁO ĐỒ ÁN TỐT NGHIỆP', 16, True)
    doc.add_paragraph()
    center_para(doc, 'ĐỀ TÀI:', 14, True)
    center_para(doc, 'XÂY DỰNG WEBSITE BÁN ĐIỆN THOẠI', 14, True)
    center_para(doc, 'VÀ PHỤ KIỆN TRỰC TUYẾN', 14, True)
    center_para(doc, 'TÍCH HỢP CHATBOT HỖ TRỢ KHÁCH HÀNG', 14, True)
    center_para(doc, '(CANDY DIGITAL)', 13, True)
    for _ in range(5):
        doc.add_paragraph()
    add_table(doc, ['', ''], [
        ['Sinh viên thực hiện:', '[Họ và tên sinh viên]'],
        ['Mã sinh viên:', '[Mã số sinh viên]'],
        ['Lớp:', '[Tên lớp]'],
        ['Giảng viên hướng dẫn:', '[Tên giảng viên]'],
    ])
    for _ in range(3):
        doc.add_paragraph()
    center_para(doc, 'TP. Hồ Chí Minh, năm 2026', 13)
    doc.add_page_break()

    # ===== NHẬN XÉT GV =====
    add_heading(doc, 'NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN', 1)
    add_para(doc, 'Họ và tên sinh viên: ........................................................................................................')
    add_para(doc, 'Mã sinh viên: ................................................  Lớp: .....................................................')
    add_para(doc, 'Tên đề tài: Xây dựng website bán điện thoại và phụ kiện trực tuyến tích hợp chatbot hỗ trợ khách hàng — Candy Digital')
    add_para(doc, '1. Nội dung và kết quả thực hiện:')
    for _ in range(6):
        add_para(doc, '............................................................................................................................................')
    add_para(doc, '2. Tinh thần, thái độ làm việc của sinh viên:')
    for _ in range(4):
        add_para(doc, '............................................................................................................................................')
    add_para(doc, '3. Nhận xét chung và đề nghị:')
    for _ in range(4):
        add_para(doc, '............................................................................................................................................')
    add_para(doc, '4. Điểm đề nghị (thang điểm 10):')
    add_para(doc, 'Điểm chấm (bằng số): ..........................   Điểm chấm (bằng chữ): .......................................')
    doc.add_paragraph()
    add_para(doc, 'TP. Hồ Chí Minh, ngày ......... tháng ......... năm 2026')
    center_para(doc, 'GIẢNG VIÊN HƯỚNG DẪN')
    center_para(doc, '(Ký và ghi rõ họ tên)')
    doc.add_page_break()

    # ===== LỜI CAM ĐOAN =====
    add_heading(doc, 'LỜI CAM ĐOAN', 1)
    add_para(doc, 'Tôi xin cam đoan đây là công trình nghiên cứu của riêng tôi. Các số liệu, kết quả nêu trong báo cáo là trung thực và chưa từng được ai công bố trong bất kỳ công trình nào khác.', indent=True)
    add_para(doc, 'Tôi xin cam đoan rằng mọi sự giúp đỡ cho việc thực hiện đề tài này đã được cảm ơn và các thông tin trích dẫn trong báo cáo đã được chỉ rõ nguồn gốc.', indent=True)
    add_para(doc, 'Nếu có bất kỳ sự gian lận nào, tôi xin chịu hoàn toàn trách nhiệm.', indent=True)
    doc.add_paragraph()
    add_para(doc, 'TP. Hồ Chí Minh, ngày ......... tháng ......... năm 2026')
    center_para(doc, 'Sinh viên thực hiện')
    center_para(doc, '(Ký và ghi rõ họ tên)')
    doc.add_page_break()

    # ===== LỜI CẢM ƠN =====
    add_heading(doc, 'LỜI CẢM ƠN', 1)
    add_para(doc, 'Lời đầu tiên, em xin gửi lời cảm ơn chân thành và sâu sắc nhất đến quý Thầy/Cô trong Khoa Công nghệ Thông tin, Trường Đại học [Tên trường] đã tận tình truyền đạt kiến thức, kinh nghiệm quý báu trong suốt quá trình học tập tại trường.', indent=True)
    add_para(doc, 'Em xin đặc biệt cảm ơn Thầy/Cô [Tên giảng viên hướng dẫn] — Giảng viên hướng dẫn đề tài — đã dành thời gian, tâm huyết hướng dẫn, góp ý và động viên em trong suốt quá trình thực hiện đề tài này. Những lời chỉ bảo của Thầy/Cô là kim chỉ nam giúp em định hướng và hoàn thiện sản phẩm.', indent=True)
    add_para(doc, 'Em cũng xin chân thành cảm ơn các bạn cùng lớp đã chia sẻ khó khăn, hỗ trợ và cùng nhau vượt qua những thử thách trong quá trình thực hiện dự án.', indent=True)
    add_para(doc, 'Mặc dù đã cố gắng hết sức, nhưng do kiến thức và kinh nghiệm còn hạn chế, báo cáo này chắc chắn không tránh khỏi những thiếu sót. Em rất mong nhận được sự góp ý, nhận xét từ quý Thầy/Cô để em có thể hoàn thiện hơn trong tương lai.', indent=True)
    add_para(doc, 'Xin trân trọng cảm ơn!', indent=True)
    doc.add_page_break()

    # ===== MỤC LỤC (placeholder) =====
    add_heading(doc, 'MỤC LỤC', 1)
    toc_items = [
        'Lời cam đoan',
        'Lời cảm ơn',
        'Mục lục',
        'Chương 1: Tổng quan đề tài',
        'Chương 2: Cơ sở lý thuyết',
        'Chương 3: Phân tích và thiết kế hệ thống',
        '  3.2 Thiết kế cơ sở dữ liệu',
        '    3.2.1 Sơ đồ quan hệ (ERD)',
        '  3.4 Sơ đồ Use Case',
        '  3.5 Sơ đồ tuần tự (Sequence)',
        '  3.6 Sơ đồ hoạt động (Activity) — Luồng đặt hàng',
        'Chương 4: Triển khai và kết quả',
        'Kết luận',
        'Tài liệu tham khảo',
        'Phụ lục: Mã nguồn chương trình',
    ]
    for item in toc_items:
        add_para(doc, item)
    add_para(doc, '(Sinh viên cập nhật mục lục tự động trong Word: References → Table of Contents)')
    doc.add_page_break()

    # ===== CHƯƠNG 1 =====
    add_heading(doc, 'CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI', 1)
    add_heading(doc, '1.1 Giới thiệu', 2)
    add_para(doc, 'Trong bối cảnh thương mại điện tử phát triển mạnh mẽ tại Việt Nam, việc xây dựng một nền tảng mua sắm trực tuyến chuyên biệt về điện thoại và phụ kiện là một hướng đi có tiềm năng cao. Đề tài "Xây dựng website bán điện thoại và phụ kiện trực tuyến tích hợp chatbot hỗ trợ khách hàng" (Candy Digital) được thực hiện nhằm đáp ứng nhu cầu mua sắm ngày càng tăng của người tiêu dùng trong lĩnh vực công nghệ.', indent=True)
    add_para(doc, 'Candy Digital là website thương mại điện tử tập trung vào các sản phẩm điện tử tiêu dùng như điện thoại thông minh, laptop, máy tính bảng, tai nghe, đồng hồ thông minh và phụ kiện đi kèm. Điểm nổi bật của hệ thống là tích hợp trợ lý tư vấn AI (Candy AI) sử dụng Google Gemini API, giúp khách hàng tra cứu sản phẩm, so sánh thông số và lựa chọn sản phẩm phù hợp nhu cầu và ngân sách mọi lúc.', indent=True)

    add_heading(doc, '1.1.1 Lý do chọn đề tài', 3)
    add_bullet(doc, 'Thị trường thương mại điện tử Việt Nam tăng trưởng trung bình 25–30% mỗi năm, đặc biệt trong lĩnh vực điện thoại và phụ kiện công nghệ.')
    add_bullet(doc, 'Nhu cầu mua sắm trực tuyến của người tiêu dùng ngày càng cao, đặc biệt sau giai đoạn dịch bệnh COVID-19.')
    add_bullet(doc, 'Công nghệ AI đang được ứng dụng rộng rãi trong chăm sóc khách hàng, tạo ra trải nghiệm mua sắm cá nhân hóa.')
    add_bullet(doc, 'Đây là cơ hội thực hành các công nghệ web hiện đại (Next.js 14, Node.js, MySQL) vào sản phẩm thực tế có đầy đủ luồng nghiệp vụ thương mại điện tử.')

    add_heading(doc, '1.2 Mục tiêu', 2)
    add_heading(doc, '1.2.1 Mục tiêu xây dựng website', 3)
    add_bullet(doc, 'Xây dựng nền tảng thương mại điện tử đầy đủ chức năng, an toàn và dễ sử dụng trên cả máy tính và thiết bị di động.')
    add_bullet(doc, 'Tích hợp hệ thống quản lý đơn hàng, giỏ hàng, thanh toán đa phương thức và khuyến mãi theo chiến dịch.')
    add_bullet(doc, 'Phát triển chatbot AI (Candy AI) hỗ trợ tư vấn sản phẩm 24/7 dựa trên catalog thực tế trong cơ sở dữ liệu.')
    add_bullet(doc, 'Xây dựng hệ thống Admin Dashboard tách biệt để quản lý sản phẩm, danh mục, đơn hàng, khuyến mãi, thanh toán và người dùng.')

    add_heading(doc, '1.2.2 Đối tượng sử dụng', 3)
    add_bullet(doc, 'Khách hàng (User): Người có nhu cầu mua điện thoại, laptop, tai nghe, phụ kiện và thiết bị công nghệ.')
    add_bullet(doc, 'Quản trị viên (Admin): Nhân viên cửa hàng quản lý toàn bộ hoạt động kinh doanh trên website.')

    add_heading(doc, '1.3 Phạm vi', 2)
    add_heading(doc, '1.3.1 Giới hạn chức năng', 3)
    add_bullet(doc, 'Khách hàng: Đăng ký/Đăng nhập, duyệt/tìm kiếm/lọc sản phẩm, chi tiết sản phẩm (màu sắc, dung lượng), giỏ hàng, đặt hàng, thanh toán (COD/chuyển khoản/MoMo/ZaloPay), xem/hủy đơn hàng, xem khuyến mãi, chatbot AI, hồ sơ cá nhân, trang giới thiệu.')
    add_bullet(doc, 'Admin: Dashboard thống kê, quản lý sản phẩm (biến thể màu/dung lượng, gallery), danh mục, chiến dịch khuyến mãi, đơn hàng (xuất Excel), người dùng, cấu hình thanh toán, tài khoản admin.')
    add_bullet(doc, 'Ngoài phạm vi: Ứng dụng di động native, tích hợp API vận chuyển (GHN/GHTK), đánh giá sản phẩm có backend, wishlist, OAuth đăng nhập mạng xã hội, thông báo email/SMS tự động.')

    add_heading(doc, '1.3.2 Công nghệ sử dụng', 3)
    add_bullet(doc, 'Frontend: Next.js 14 (App Router), React 18, TypeScript, Tailwind CSS, Zustand, Axios.')
    add_bullet(doc, 'Backend: Node.js, Express.js 4, JWT, Bcryptjs, Multer, ExcelJS, express-validator.')
    add_bullet(doc, 'Database: MySQL 8.0 (mysql2 driver, truy vấn SQL trực tiếp).')
    add_bullet(doc, 'AI Chatbot: Google Gemini API (model gemini-2.5-flash).')
    add_bullet(doc, 'Dịch vụ bổ trợ: VietQR (mã QR chuyển khoản), webhook sao kê ngân hàng (Sepay).')

    add_heading(doc, '1.4 Phương pháp thực hiện', 2)
    add_heading(doc, '1.4.1 Nghiên cứu tài liệu', 3)
    add_bullet(doc, 'Nghiên cứu tài liệu chính thức của Next.js 14, Express.js, MySQL 8.0 và Google Gemini API.')
    add_bullet(doc, 'Phân tích các website thương mại điện tử lớn như Shopee, Tiki, Thế Giới Di Động.')
    add_bullet(doc, 'Tìm hiểu design pattern: MVC, RESTful API, JWT Authentication, Database Transaction.')

    add_heading(doc, '1.4.2 Thiết kế và lập trình', 3)
    add_bullet(doc, 'Phân tích yêu cầu, xác định chức năng cho từng vai trò người dùng (Guest, Customer, Admin).')
    add_bullet(doc, 'Thiết kế ERD, xác định bảng dữ liệu và mối quan hệ.')
    add_bullet(doc, 'Lập trình theo quy trình: Database Schema → Backend API → Frontend UI.')
    add_bullet(doc, 'Kiểm thử từng chức năng với Postman (API) và trình duyệt (UI).')
    doc.add_page_break()

    # ===== CHƯƠNG 2 =====
    add_heading(doc, 'CHƯƠNG 2: CƠ SỞ LÝ THUYẾT', 1)
    add_heading(doc, '2.1 Tổng quan về Website Thương Mại Điện Tử', 2)
    add_heading(doc, '2.1.1 Khái niệm', 3)
    add_para(doc, 'Website thương mại điện tử là nền tảng trực tuyến cho phép thực hiện các giao dịch mua bán qua internet, bao gồm hệ thống giỏ hàng, quản lý đơn hàng, xác thực người dùng và bảo mật dữ liệu giao dịch.', indent=True)
    add_heading(doc, '2.1.2 Phân loại', 3)
    add_bullet(doc, 'B2C (Business to Consumer): Bán hàng trực tiếp đến người tiêu dùng — mô hình của Candy Digital.')
    add_bullet(doc, 'B2B (Business to Business): Giao dịch giữa các doanh nghiệp.')
    add_bullet(doc, 'C2C (Consumer to Consumer): Giao dịch giữa người dùng như Chợ Tốt.')

    add_heading(doc, '2.2 Công nghệ sử dụng', 2)
    add_heading(doc, '2.2.1 Next.js 14 — Frontend', 3)
    add_para(doc, 'Next.js là framework React hỗ trợ SSR/SSG và App Router với React Server Components. Phiên bản 14 cải thiện hiệu suất đáng kể, hỗ trợ TypeScript native và tối ưu hóa hình ảnh tự động. Dự án sử dụng kiến trúc dual-app: cùng một codebase chạy hai instance độc lập — cổng 3000 cho khách hàng và cổng 3001 cho quản trị viên.', indent=True)

    add_heading(doc, '2.2.2 Node.js & Express.js — Backend', 3)
    add_para(doc, 'Node.js runtime với kiến trúc event-driven non-blocking I/O. Express.js cung cấp hệ thống routing và middleware linh hoạt để xây dựng RESTful API với JWT Authentication, Multer xử lý upload file và ExcelJS xuất báo cáo đơn hàng.', indent=True)

    add_heading(doc, '2.2.3 MySQL 8.0 — Cơ sở dữ liệu', 3)
    add_para(doc, 'RDBMS phổ biến, hỗ trợ ACID transaction, Foreign Key Constraint và đảm bảo tính toàn vẹn dữ liệu. Hệ thống sử dụng Transaction trong quy trình đặt hàng để đảm bảo tính nguyên tử: tạo đơn hàng → tạo chi tiết đơn → trừ tồn kho → xóa giỏ hàng.', indent=True)

    add_heading(doc, '2.2.4 Tailwind CSS', 3)
    add_para(doc, 'Framework CSS utility-first, xây dựng giao diện nhanh chóng qua class tiện ích trực tiếp trong JSX, không cần viết file CSS riêng, code nhất quán và dễ bảo trì.', indent=True)

    add_heading(doc, '2.2.5 Google Gemini API — Trí tuệ nhân tạo', 3)
    add_para(doc, 'Google Gemini là mô hình ngôn ngữ lớn (LLM) của Google, cung cấp khả năng xử lý ngôn ngữ tự nhiên (NLP) tiên tiến. Dự án sử dụng model gemini-2.5-flash với system prompt tùy chỉnh đưa catalog 30 sản phẩm đang bán vào context, lưu lịch sử hội thoại vào database để duy trì ngữ cảnh phiên chat.', indent=True)

    add_heading(doc, '2.2.6 Zustand — State Management', 3)
    add_para(doc, 'Thư viện quản lý state nhẹ nhàng cho React, quản lý trạng thái đăng nhập (authStore) và giỏ hàng (cartStore), đồng bộ trên toàn bộ ứng dụng khách hàng.', indent=True)

    add_heading(doc, '2.3 Công cụ hỗ trợ', 2)
    add_bullet(doc, 'Visual Studio Code: Trình soạn thảo code với extension TypeScript, ESLint, Tailwind CSS IntelliSense.')
    add_bullet(doc, 'Git & GitHub: Hệ thống kiểm soát phiên bản, quản lý lịch sử thay đổi code.')
    add_bullet(doc, 'Postman: Công cụ kiểm thử API REST.')
    add_bullet(doc, 'MySQL Workbench: Công cụ quản lý CSDL trực quan, thiết kế ERD, chạy SQL script.')
    doc.add_page_break()

    # ===== CHƯƠNG 3 =====
    add_heading(doc, 'CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG', 1)
    add_heading(doc, '3.1 Phân tích yêu cầu', 2)
    add_heading(doc, '3.1.1 Yêu cầu chức năng', 3)
    add_para(doc, 'Hệ thống phục vụ ba nhóm người dùng: Khách vãng lai (Guest), Khách hàng đã đăng nhập (Customer) và Quản trị viên (Admin).')
    add_para(doc, 'Bảng 3.1: Chức năng dành cho Khách hàng (User)', bold=True)
    add_table(doc, ['STT', 'Chức năng', 'Mô tả'], [
        ['1', 'Đăng ký / Đăng nhập', 'Tạo tài khoản bằng email/mật khẩu; đăng nhập xác thực JWT; đăng xuất xóa token.'],
        ['2', 'Xem sản phẩm', 'Duyệt danh sách theo danh mục cha/con, tìm kiếm tên/thương hiệu, sắp xếp (mới nhất, giá, tên), phân trang.'],
        ['3', 'Chi tiết sản phẩm', 'Gallery ảnh theo màu sắc, chọn dung lượng (giá/tồn kho riêng), mô tả, thông số kỹ thuật, highlights, sản phẩm liên quan.'],
        ['4', 'Khuyến mãi', 'Xem chiến dịch đang chạy (11.11, 12.12, Black Friday…), countdown, giá giảm realtime.'],
        ['5', 'Giỏ hàng', 'Thêm/sửa/xóa số lượng, snapshot giá khuyến mãi tại thời điểm thêm, xem tổng tiền.'],
        ['6', 'Đặt hàng', 'Nhập địa chỉ giao hàng, chọn PT thanh toán (COD/Bank/MoMo/ZaloPay), ghi chú, hiển thị QR thanh toán.'],
        ['7', 'Xem đơn hàng', 'Xem lịch sử đơn, lọc trạng thái, theo dõi thanh toán, hủy đơn (pending/confirmed).'],
        ['8', 'Chat AI (Candy AI)', 'Tư vấn sản phẩm 24/7 qua chatbot Gemini AI dựa trên catalog thực tế; câu hỏi nhanh; xóa lịch sử.'],
        ['9', 'Hồ sơ cá nhân', 'Cập nhật thông tin (tên, SĐT, địa chỉ), đổi mật khẩu.'],
        ['10', 'Giới thiệu', 'Trang About — tầm nhìn, sứ mệnh, lịch sử phát triển Candy Digital.'],
    ])

    add_para(doc, 'Bảng 3.2: Chức năng dành cho Quản trị viên (Admin)', bold=True)
    add_table(doc, ['STT', 'Chức năng', 'Mô tả'], [
        ['1', 'Dashboard', 'KPI doanh thu tháng, số đơn, khách hàng, sản phẩm; biểu đồ doanh thu 6 tháng; 5 đơn gần nhất.'],
        ['2', 'Quản lý sản phẩm', 'CRUD sản phẩm; upload gallery (≤10 ảnh); quản lý màu sắc + ảnh màu; dung lượng (giá/tồn kho); highlights; thông số kỹ thuật; bật/tắt.'],
        ['3', 'Quản lý danh mục', 'CRUD danh mục cha/con; upload ảnh; bật/tắt; không xóa nếu còn sản phẩm.'],
        ['4', 'Quản lý khuyến mãi', 'CRUD chiến dịch; preset 11.11/12.12/Black Friday; phạm vi all/category/product; priority; bật/tắt nhanh.'],
        ['5', 'Quản lý đơn hàng', 'Danh sách + lọc/tìm kiếm/khoảng ngày; cập nhật trạng thái; xác nhận thanh toán; xuất Excel 2 sheet.'],
        ['6', 'Quản lý người dùng', 'Danh sách khách; tìm kiếm; khóa/mở khóa tài khoản.'],
        ['7', 'Cấu hình thanh toán', 'Bật/tắt COD/Bank/MoMo/ZaloPay; STK, tên TK, mã ngân hàng; upload QR; preview VietQR.'],
        ['8', 'Quản lý tài khoản Admin', 'Tạo/xóa admin; không xóa chính mình.'],
    ])

    add_heading(doc, '3.1.2 Yêu cầu phi chức năng', 3)
    add_bullet(doc, 'Hiệu suất: Tải trang < 3 giây; API response < 500ms trong môi trường local.')
    add_bullet(doc, 'Bảo mật: Bcrypt hash mật khẩu (salt=10); JWT thời hạn 7 ngày; middleware phân quyền user/admin.')
    add_bullet(doc, 'Khả năng mở rộng: Kiến trúc tách biệt frontend/backend, RESTful API, schema auto-migration.')
    add_bullet(doc, 'Giao diện: Responsive từ 375px đến 1920px, thiết kế hiện đại nhất quán trên cả hai ứng dụng.')

    add_heading(doc, '3.2 Thiết kế cơ sở dữ liệu', 2)
    add_heading(doc, '3.2.1 Sơ đồ quan hệ (ERD)', 3)
    add_para(doc, 'Hệ thống sử dụng 14 bảng dữ liệu kết nối qua khóa ngoại. Các quan hệ chính:', indent=True)
    add_bullet(doc, 'users (1) → (N) orders, carts, chat_sessions')
    add_bullet(doc, 'categories (1) → (N) products; categories self-referencing (parent_id)')
    add_bullet(doc, 'products (1) → (N) product_images, product_colors, product_variants')
    add_bullet(doc, 'orders (1) → (N) order_items')
    add_bullet(doc, 'chat_sessions (1) → (N) chat_messages')
    add_bullet(doc, 'promotion_campaigns (1) → (N) promotion_campaign_targets')
    add_bullet(doc, 'payment_settings: bảng cấu hình độc lập cho phương thức thanh toán')
    add_para(doc, 'Sơ đồ ERD dưới đây mô tả trực quan mối quan hệ giữa 14 bảng dữ liệu, bao gồm khóa chính (PK), khóa ngoại (FK) và cardinality (1:N).', indent=True)
    add_image(doc, os.path.join(DIAGRAM_DIR, 'erd_diagram.png'),
              'Hình 3.1: Sơ đồ quan hệ thực thể (ERD) — Cơ sở dữ liệu Candy Digital', width=Inches(6.8))
    add_para(doc, 'Hình 3.1 thể hiện nhóm bảng sản phẩm (products, product_images, product_colors, product_variants) liên kết với categories; nhóm giao dịch (carts, orders, order_items) liên kết với users; nhóm chatbot (chat_sessions, chat_messages) và nhóm khuyến mãi (promotion_campaigns, promotion_campaign_targets) tách biệt rõ ràng.', indent=True)

    add_heading(doc, '3.2.2 Tổng quan các bảng dữ liệu', 3)
    add_para(doc, 'Bảng 3.3: Tổng quan cấu trúc cơ sở dữ liệu', bold=True)
    add_table(doc, ['STT', 'Tên bảng', 'Chức năng'], [
        ['1', 'users', 'Tài khoản người dùng và quản trị viên'],
        ['2', 'categories', 'Danh mục sản phẩm, hỗ trợ cha/con'],
        ['3', 'products', 'Thông tin sản phẩm: tên, giá, tồn kho, thương hiệu, mô tả'],
        ['4', 'product_images', 'Hình ảnh sản phẩm, nhiều ảnh/sản phẩm'],
        ['5', 'product_colors', 'Màu sắc sản phẩm, mã hex, gallery theo màu'],
        ['6', 'product_variants', 'Biến thể dung lượng: giá và tồn kho riêng'],
        ['7', 'carts', 'Giỏ hàng (kèm variant, màu, giá snapshot khuyến mãi)'],
        ['8', 'orders', 'Đơn hàng: giao hàng, thanh toán, trạng thái, webhook ngân hàng'],
        ['9', 'order_items', 'Chi tiết sản phẩm trong đơn (snapshot tại thời điểm đặt)'],
        ['10', 'chat_sessions', 'Phiên chat AI, hỗ trợ khách vãng lai'],
        ['11', 'chat_messages', 'Lịch sử tin nhắn người dùng và AI'],
        ['12', 'payment_settings', 'Cấu hình phương thức thanh toán shop'],
        ['13', 'promotion_campaigns', 'Chiến dịch khuyến mãi (% giảm, thời gian, phạm vi)'],
        ['14', 'promotion_campaign_targets', 'Đối tượng áp dụng khuyến mãi (category/product)'],
    ])

    add_heading(doc, '3.2.3 Chi tiết một số bảng quan trọng', 3)
    add_para(doc, 'Bảng 3.4: Chi tiết bảng users', bold=True)
    add_table(doc, ['Tên trường', 'Kiểu dữ liệu', 'Ràng buộc', 'Ý nghĩa'], [
        ['id', 'INT', 'PK, AUTO_INCREMENT', 'Mã định danh người dùng'],
        ['name', 'VARCHAR(100)', 'NOT NULL', 'Họ và tên'],
        ['email', 'VARCHAR(100)', 'NOT NULL, UNIQUE', 'Email đăng nhập'],
        ['password', 'VARCHAR(255)', 'NOT NULL', 'Mật khẩu bcrypt'],
        ['phone', 'VARCHAR(15)', 'NULL', 'Số điện thoại'],
        ['address', 'TEXT', 'NULL', 'Địa chỉ mặc định'],
        ['role', 'ENUM', "DEFAULT 'user'", 'user hoặc admin'],
        ['is_active', 'TINYINT(1)', 'DEFAULT 1', '1=hoạt động, 0=khóa'],
        ['created_at', 'TIMESTAMP', 'DEFAULT NOW()', 'Thời điểm tạo'],
    ])

    add_para(doc, 'Bảng 3.5: Chi tiết bảng orders', bold=True)
    add_table(doc, ['Tên trường', 'Kiểu dữ liệu', 'Ràng buộc', 'Ý nghĩa'], [
        ['id', 'INT', 'PK', 'Mã đơn hàng'],
        ['user_id', 'INT', 'FK → users', 'Người đặt hàng'],
        ['total_price', 'DECIMAL(15,2)', 'NOT NULL', 'Tổng tiền (VNĐ)'],
        ['shipping_name/phone/address', 'VARCHAR/TEXT', 'NOT NULL', 'Thông tin giao hàng'],
        ['payment_method', 'ENUM', "DEFAULT 'cod'", 'cod/bank_transfer/momo/zalopay'],
        ['payment_status', 'VARCHAR(20)', "DEFAULT 'na'", 'na/pending/paid'],
        ['status', 'ENUM', "DEFAULT 'pending'", 'pending/confirmed/shipping/done/cancelled'],
        ['bank_in_*', 'MIXED', 'NULL', 'Dữ liệu webhook sao kê ngân hàng'],
        ['note', 'TEXT', 'NULL', 'Ghi chú khách hàng'],
    ])

    add_heading(doc, '3.3 Thiết kế kiến trúc hệ thống', 2)
    add_para(doc, 'Hệ thống theo kiến trúc Client-Server ba tầng:', indent=True)
    add_bullet(doc, 'Tầng trình diễn: Next.js 14 — ứng dụng khách hàng (cổng 3000) và quản trị (cổng 3001/admin), cùng codebase, tách biệt localStorage/token.')
    add_bullet(doc, 'Tầng xử lý: Express.js REST API (cổng 5000) — 40+ endpoint, middleware JWT, upload Multer.')
    add_bullet(doc, 'Tầng dữ liệu: MySQL 8.0 — 14 bảng, transaction khi đặt hàng, schema auto-migration khi khởi động server.')
    add_bullet(doc, 'Dịch vụ ngoài: Google Gemini API (chatbot), VietQR (QR chuyển khoản), webhook Sepay (xác nhận TT tự động).')

    add_heading(doc, '3.4 Sơ đồ Use Case', 2)
    add_para(doc, 'Sơ đồ Use Case mô tả tương tác giữa các tác nhân (Actor) và hệ thống Candy Digital. Hệ thống có ba tác nhân chính: Khách vãng lai (chỉ xem sản phẩm và chat AI), Khách hàng (đã đăng nhập, thực hiện mua hàng) và Quản trị viên (quản lý toàn bộ hệ thống).', indent=True)
    add_para(doc, 'Bảng 3.6: Mô tả các tác nhân hệ thống', bold=True)
    add_table(doc, ['Tác nhân', 'Mô tả', 'Quyền hạn chính'], [
        ['Khách vãng lai', 'Người truy cập chưa đăng nhập', 'Xem SP, tìm kiếm, chat AI, xem khuyến mãi'],
        ['Khách hàng', 'Người dùng đã đăng ký tài khoản', 'Mua hàng, giỏ hàng, đặt hàng, quản lý đơn, hồ sơ'],
        ['Quản trị viên', 'Nhân viên quản lý cửa hàng', 'CRUD SP/DM/KM, quản lý đơn, cấu hình TT, thống kê'],
    ])
    add_image(doc, os.path.join(DIAGRAM_DIR, 'use_case_overall.png'),
              'Hình 3.2: Sơ đồ Use Case tổng quát hệ thống Candy Digital')
    add_para(doc, 'Hình 3.2 thể hiện toàn bộ use case của hệ thống trong một biên hệ thống (system boundary). Khách vãng lai và khách hàng chia sẻ các use case xem sản phẩm, tìm kiếm và chat AI; riêng khách hàng mới có thể đặt hàng và quản lý giỏ hàng.', indent=True)
    add_image(doc, os.path.join(DIAGRAM_DIR, 'use_case_customer.png'),
              'Hình 3.3: Sơ đồ Use Case — Module Khách hàng')
    add_para(doc, 'Hình 3.3 chi tiết hóa các use case phía khách hàng, bao gồm luồng chọn biến thể sản phẩm (màu sắc, dung lượng), chọn phương thức thanh toán và tương tác với chatbot Candy AI.', indent=True)
    add_image(doc, os.path.join(DIAGRAM_DIR, 'use_case_admin.png'),
              'Hình 3.4: Sơ đồ Use Case — Module Quản trị viên')
    add_para(doc, 'Hình 3.4 mô tả các use case quản trị: quản lý sản phẩm/danh mục/khuyến mãi, xử lý đơn hàng, xác nhận thanh toán, xuất báo cáo Excel và cấu hình thông tin thanh toán shop.', indent=True)

    add_heading(doc, '3.5 Sơ đồ tuần tự (Sequence Diagram)', 2)
    add_para(doc, 'Sơ đồ tuần tự mô tả trình tự trao đổi message giữa các thành phần theo thời gian. Dưới đây là bốn luồng nghiệp vụ quan trọng nhất của hệ thống.', indent=True)

    add_heading(doc, '3.5.1 Luồng đăng nhập', 3)
    add_para(doc, 'Khi khách hàng đăng nhập, Frontend gửi email/mật khẩu đến Backend. Backend truy vấn MySQL, so sánh mật khẩu bằng bcrypt, tạo JWT token (thời hạn 7 ngày) và trả về cho Frontend lưu vào localStorage.', indent=True)
    add_image(doc, os.path.join(DIAGRAM_DIR, 'seq_login.png'),
              'Hình 3.5: Sơ đồ tuần tự — Đăng nhập hệ thống')

    add_heading(doc, '3.5.2 Luồng đặt hàng (Checkout)', 3)
    add_para(doc, 'Luồng đặt hàng sử dụng MySQL Transaction để đảm bảo tính nguyên tử (ACID). Các bước: kiểm tra giỏ hàng và tồn kho → tạo bản ghi orders và order_items → trừ tồn kho (products hoặc product_variants) → xóa giỏ hàng → commit. Nếu bất kỳ bước nào thất bại, toàn bộ transaction được rollback.', indent=True)
    add_image(doc, os.path.join(DIAGRAM_DIR, 'seq_order.png'),
              'Hình 3.6: Sơ đồ tuần tự — Đặt hàng (Checkout)')

    add_heading(doc, '3.5.3 Luồng Chatbot AI (Candy AI)', 3)
    add_para(doc, 'Chatbot hoạt động theo mô hình session-based. Frontend lấy hoặc tạo session_token, lưu vào localStorage. Mỗi tin nhắn được Backend lưu vào chat_messages, truy vấn 30 sản phẩm active làm context, gọi Google Gemini API với system prompt tùy chỉnh, rồi lưu và trả về phản hồi AI.', indent=True)
    add_image(doc, os.path.join(DIAGRAM_DIR, 'seq_chatbot.png'),
              'Hình 3.7: Sơ đồ tuần tự — Chatbot AI (Candy AI)')

    add_heading(doc, '3.5.4 Luồng xác nhận thanh toán chuyển khoản', 3)
    add_para(doc, 'Khi khách chuyển khoản với nội dung chứa mã đơn (DHxxxxxx), dịch vụ sao kê (Sepay) gửi webhook đến Backend. Backend đối chiếu số tiền và mã đơn, cập nhật payment_status = paid nếu khớp.', indent=True)
    add_image(doc, os.path.join(DIAGRAM_DIR, 'seq_payment.png'),
              'Hình 3.8: Sơ đồ tuần tự — Xác nhận thanh toán chuyển khoản')

    add_para(doc, 'Bảng 3.7: Tóm tắt các luồng tuần tự', bold=True)
    add_table(doc, ['Luồng', 'API chính', 'Công nghệ liên quan'], [
        ['Đăng nhập', 'POST /api/auth/login', 'JWT, bcrypt, MySQL'],
        ['Đặt hàng', 'POST /api/orders', 'MySQL Transaction, stock update'],
        ['Chatbot', 'GET/POST /api/chat/*', 'Google Gemini API, session token'],
        ['Thanh toán CK', 'POST /api/orders/webhook/bank-incoming', 'Webhook Sepay, VietQR'],
    ])

    add_heading(doc, '3.6 Sơ đồ hoạt động (Activity Diagram) — Luồng đặt hàng', 2)
    add_para(doc, 'Sơ đồ hoạt động (Activity Diagram) mô tả luồng xử lý nghiệp vụ đặt hàng từ khi khách chọn sản phẩm đến khi hoàn tất đơn hàng. Sơ đồ sử dụng các ký hiệu UML chuẩn: hình tròn đặc (Start/End), hình chữ nhật bo góc (Activity), hình thoi (Decision) và mũi tên có nhãn (luồng điều kiện).', indent=True)
    add_para(doc, 'Bảng 3.8: Mô tả các bước trong luồng đặt hàng', bold=True)
    add_table(doc, ['Bước', 'Hoạt động', 'Mô tả'], [
        ['1', 'Chọn sản phẩm', 'Khách chọn màu sắc, dung lượng trên trang chi tiết SP'],
        ['2', 'Thêm giỏ hàng', 'Backend snapshot giá khuyến mãi vào carts.unit_price'],
        ['3', 'Xác thực', 'Kiểm tra đăng nhập; nếu chưa → chuyển trang login'],
        ['4', 'Checkout', 'Nhập địa chỉ giao hàng, chọn PT thanh toán, ghi chú'],
        ['5', 'Transaction', 'BEGIN → kiểm tra tồn kho → INSERT order → trừ stock → xóa cart → COMMIT'],
        ['6', 'Kết quả', 'COD: xác nhận đơn | Trả trước: hiển thị QR/thông tin chuyển khoản'],
    ])
    add_image(doc, os.path.join(DIAGRAM_DIR, 'activity_order.png'),
              'Hình 3.9: Sơ đồ hoạt động — Luồng đặt hàng (Checkout)', width=Inches(5.2))
    add_para(doc, 'Hình 3.9 minh họa điểm quyết định quan trọng: (1) kiểm tra đăng nhập trước checkout; (2) kiểm tra tồn kho trong transaction — nếu không đủ thì ROLLBACK và thông báo lỗi; (3) phân nhánh theo phương thức thanh toán sau khi COMMIT thành công. Cơ chế transaction đảm bảo không xảy ra tình trạng trừ tồn kho mà không tạo được đơn hàng.', indent=True)
    doc.add_page_break()

    # ===== CHƯƠNG 4 =====
    add_heading(doc, 'CHƯƠNG 4: TRIỂN KHAI VÀ KẾT QUẢ', 1)
    add_heading(doc, '4.1 Xây dựng hệ thống', 2)
    add_heading(doc, '4.1.1 Cấu trúc dự án', 3)
    add_para(doc, 'Dự án được tổ chức theo cấu trúc monorepo gồm thư mục database/, backend/ và frontend/. Backend chia theo controllers, routes, middlewares, config. Frontend sử dụng Next.js App Router với thư mục app/, components/, store/ (Zustand), lib/api.ts (Axios).', indent=True)

    add_heading(doc, '4.1.2 Các luồng nghiệp vụ chính', 3)
    add_para(doc, 'Luồng đặt hàng: Khách chọn sản phẩm (màu, dung lượng) → thêm giỏ (snapshot giá KM) → checkout (nhập địa chỉ, chọn PT TT) → backend transaction (tạo order + order_items + trừ stock + xóa cart) → hiển thị QR/thông tin TT nếu trả trước.', indent=True)
    add_para(doc, 'Luồng khuyến mãi: Admin tạo campaign (% giảm, thời gian, phạm vi) → backend tính sale_price realtime → giá khóa vào cart.unit_price khi thêm giỏ.', indent=True)
    add_para(doc, 'Luồng chatbot: Frontend lấy/tạo session token → gửi tin nhắn → backend inject catalog 30 SP vào system prompt → gọi Gemini API → lưu reply vào chat_messages.', indent=True)
    add_para(doc, 'Luồng thanh toán chuyển khoản: Sau đặt hàng hiển thị QR VietQR + STK shop → webhook Sepay khớp mã DHxxxxxx → tự xác nhận payment_status=paid (nếu cấu hình).', indent=True)

    add_heading(doc, '4.2 Kết quả đạt được', 2)
    add_para(doc, 'Hệ thống đã hoàn thiện đầy đủ các chức năng đề ra. Dưới đây là mô tả giao diện từng module (sinh viên chèn ảnh chụp màn hình tương ứng):', indent=True)

    add_heading(doc, '4.2.1 Giao diện phía Khách hàng', 3)
    sections_customer = [
        ('a) Trang chủ', 'Hero carousel 4 slide, flash sale, danh mục icon, grid sản phẩm nổi bật, banner chính sách (miễn phí vận chuyển, bảo hành).'),
        ('b) Trang đăng nhập / đăng ký', 'Form xác thực JWT, nút ẩn/hiện mật khẩu, chuyển trang đăng ký/đăng nhập.'),
        ('c) Danh sách sản phẩm', 'Grid responsive, sidebar lọc danh mục, tìm kiếm, sắp xếp, phân trang, hiển thị giá khuyến mãi.'),
        ('d) Chi tiết sản phẩm', 'Gallery theo màu, chọn dung lượng, accordion thông số kỹ thuật, highlights, mua ngay/thêm giỏ, sản phẩm liên quan.'),
        ('e) Trang khuyến mãi', 'Campaign đang chạy, countdown, danh sách sản phẩm giảm giá.'),
        ('f) Giỏ hàng', 'Danh sách SP, điều chỉnh số lượng, xóa, tổng tiền, nút checkout.'),
        ('g) Checkout', 'Form giao hàng, chọn PT thanh toán (COD/Bank/MoMo/ZaloPay), tóm tắt đơn, hiển thị QR sau đặt hàng.'),
        ('h) Đơn hàng của tôi', 'Tabs lọc trạng thái, chi tiết đơn, hủy đơn, theo dõi thanh toán.'),
        ('i) Hồ sơ cá nhân', 'Cập nhật thông tin, đổi mật khẩu.'),
        ('j) Chat AI (Candy AI)', 'Widget floating góc phải, câu hỏi nhanh, lịch sử chat, xóa lịch sử.'),
        ('k) Trang giới thiệu', 'Tầm nhìn, sứ mệnh, timeline phát triển Candy Digital.'),
    ]
    fig = 1
    for title, desc in sections_customer:
        add_para(doc, title, bold=True)
        add_para(doc, desc, indent=True)
        add_para(doc, f'Hình 4.{fig}: Giao diện {title[3:]}')
        fig += 1

    add_heading(doc, '4.2.2 Giao diện phía Quản trị viên', 3)
    sections_admin = [
        ('a) Admin Dashboard', '4 KPI cards, biểu đồ doanh thu 6 tháng, bảng 5 đơn gần nhất.'),
        ('b) Quản lý sản phẩm', 'Bảng SP + form CRUD, upload gallery, quản lý màu/dung lượng/thông số.'),
        ('c) Quản lý danh mục', 'Cấu trúc cha/con, CRUD, upload ảnh.'),
        ('d) Quản lý khuyến mãi', 'CRUD campaign, preset 11.11/12.12/Black Friday, bật/tắt nhanh.'),
        ('e) Quản lý đơn hàng', 'Lọc/tìm kiếm, cập nhật trạng thái, xác nhận TT, xuất Excel.'),
        ('f) Quản lý người dùng', 'Danh sách, tìm kiếm, khóa/mở khóa.'),
        ('g) Cấu hình thanh toán', 'Bật/tắt PT TT, STK, QR, preview VietQR.'),
        ('h) Quản lý tài khoản Admin', 'Tạo/xóa admin.'),
    ]
    for title, desc in sections_admin:
        add_para(doc, title, bold=True)
        add_para(doc, desc, indent=True)
        add_para(doc, f'Hình 4.{fig}: Giao diện {title[3:]}')
        fig += 1

    add_heading(doc, '4.3 Đánh giá', 2)
    add_heading(doc, '4.3.1 Ưu điểm', 3)
    add_bullet(doc, 'Kiến trúc hiện đại: Next.js 14 App Router, TypeScript, dual-app pattern, Zustand theo best practice.')
    add_bullet(doc, 'Bảo mật tốt: JWT, bcrypt hash, middleware phân quyền, MySQL transaction khi đặt hàng.')
    add_bullet(doc, 'UX mượt mà: Loading states, toast notification, responsive design, giao diện chi tiết SP kiểu TGDD.')
    add_bullet(doc, 'AI độc đáo: Chatbot Candy AI tư vấn dựa trên catalog thực tế, không cần đăng nhập.')
    add_bullet(doc, 'Khuyến mãi linh hoạt: Campaign theo phạm vi/priority, snapshot giá giỏ hàng.')
    add_bullet(doc, 'Thanh toán đa dạng: COD + chuyển khoản (VietQR) + MoMo/ZaloPay + webhook tự xác nhận.')
    add_bullet(doc, 'Admin đầy đủ: Dashboard, xuất Excel, cấu hình thanh toán linh hoạt.')

    add_heading(doc, '4.3.2 Nhược điểm', 3)
    add_bullet(doc, 'MoMo/ZaloPay chưa tích hợp SDK chính thức — hiển thị QR/thông tin chuyển khoản thủ công.')
    add_bullet(doc, 'Chưa có hệ thống email/SMS tự động (xác nhận đơn, quên mật khẩu).')
    add_bullet(doc, 'Đánh giá sản phẩm chỉ có UI placeholder, chưa có backend.')
    add_bullet(doc, 'Lưu ảnh cục bộ (backend/uploads/), chưa dùng cloud storage.')
    add_bullet(doc, 'Chưa có automated test (unit test, integration test).')
    add_bullet(doc, 'Chưa tích hợp API vận chuyển thực tế (GHN, GHTK).')

    add_heading(doc, '4.4 Hướng phát triển', 2)
    add_bullet(doc, 'Tích hợp SDK thanh toán MoMo, ZaloPay, VNPay chính thức.')
    add_bullet(doc, 'Hệ thống email tự động với NodeMailer/SendGrid.')
    add_bullet(doc, 'Backend đánh giá sản phẩm, wishlist, mã coupon.')
    add_bullet(doc, 'Phát triển ứng dụng di động React Native.')
    add_bullet(doc, 'Redis cache, CDN hình ảnh (Cloudinary/S3).')
    add_bullet(doc, 'Tích hợp API vận chuyển và tracking đơn hàng.')
    add_bullet(doc, 'CI/CD Pipeline với GitHub Actions, Docker.')
    doc.add_page_break()

    # ===== KẾT LUẬN =====
    add_heading(doc, 'KẾT LUẬN', 1)
    add_para(doc, 'Qua quá trình thực hiện đề tài, em đã hoàn thành hệ thống Candy Digital với kiến trúc kỹ thuật hiện đại: Next.js 14 frontend (dual-app khách hàng/admin), Express.js/Node.js backend với hơn 40 REST endpoint, MySQL 8.0 với 14 bảng được chuẩn hóa và chatbot AI Candy AI tích hợp Google Gemini API.', indent=True)
    add_para(doc, 'Website đã triển khai đầy đủ 11 trang khách hàng và 8 trang quản trị viên, đáp ứng tất cả yêu cầu chức năng đề ra. Tính năng chatbot AI tư vấn sản phẩm và hệ thống khuyến mãi theo chiến dịch là những điểm nổi bật, mang lại trải nghiệm mua sắm thông minh và hiện đại.', indent=True)
    add_para(doc, 'Qua đề tài này, em đã củng cố kỹ năng lập trình web fullstack, thiết kế CSDL quan hệ, xây dựng RESTful API chuẩn, tích hợp dịch vụ AI và triển khai luồng thanh toán đa phương thức — nền tảng vững chắc cho sự nghiệp phát triển phần mềm.', indent=True)
    add_para(doc, 'Mặc dù còn một số hạn chế đã được nêu rõ, các chức năng cốt lõi đã hoạt động ổn định và hướng phát triển tương lai đã được vạch rõ. Candy Digital hoàn toàn có tiềm năng trở thành sản phẩm thương mại điện tử hoàn chỉnh.', indent=True)
    doc.add_page_break()

    # ===== TÀI LIỆU THAM KHẢO =====
    add_heading(doc, 'TÀI LIỆU THAM KHẢO', 1)
    add_heading(doc, 'Tài liệu chính thức', 3)
    refs = [
        'Next.js Documentation (2024). App Router, Server Components. https://nextjs.org/docs',
        'Express.js Documentation (2024). Routing, Middleware. https://expressjs.com/',
        'MySQL 8.0 Reference Manual (2024). https://dev.mysql.com/doc/',
        'Google AI Gemini API Documentation (2024). https://ai.google.dev/gemini-api/docs',
        'Tailwind CSS Documentation (2024). https://tailwindcss.com/docs',
        'Zustand Documentation (2024). https://zustand-demo.pmnd.rs/',
        'VietQR API (2024). https://vietqr.io/',
    ]
    for r in refs:
        add_bullet(doc, r)

    add_heading(doc, 'Sách tham khảo', 3)
    books = [
        'Marijn Haverbeke (2018). Eloquent JavaScript, 3rd Edition. No Starch Press.',
        'Mario Casciaro & Luciano Mammino (2021). Node.js Design Patterns, 3rd Edition. Packt.',
        'Michael J. Hernandez (2020). Database Design for Mere Mortals, 4th Ed. Addison-Wesley.',
    ]
    for b in books:
        add_bullet(doc, b)

    add_heading(doc, 'Tài liệu trực tuyến', 3)
    online = [
        'MDN Web Docs (2024). JavaScript Reference. https://developer.mozilla.org/',
        'Bộ Công Thương VN (2023). Sách trắng TMĐT Việt Nam 2023. https://moit.gov.vn/',
        'Statista (2024). E-commerce revenue Vietnam 2023–2028. https://www.statista.com/',
        'JWT.io — Introduction to JSON Web Tokens. https://jwt.io/introduction',
    ]
    for o in online:
        add_bullet(doc, o)
    doc.add_page_break()

    # ===== PHỤ LỤC =====
    add_heading(doc, 'PHỤ LỤC: MÃ NGUỒN CHƯƠNG TRÌNH', 1)
    add_para(doc, 'Phụ lục trình bày các đoạn mã nguồn tiêu biểu của hệ thống Candy Digital, được trích từ dự án thực tế. Mã nguồn đầy đủ được lưu tại thư mục dự án candy-digital/.', indent=True)

    add_heading(doc, 'Phụ lục A: Xác thực người dùng (Đăng ký & Đăng nhập)', 2)
    add_para(doc, 'Module auth.controller.js xử lý đăng ký và đăng nhập. Mật khẩu được mã hóa bằng bcrypt (salt rounds = 10), token JWT có thời hạn 7 ngày.', indent=True)
    add_code_block(doc, 'Listing A.1 — Hàm đăng ký và đăng nhập',
                   read_source('backend/src/controllers/auth.controller.js', 1, 75),
                   'backend/src/controllers/auth.controller.js')

    add_heading(doc, 'Phụ lục B: Đặt hàng với MySQL Transaction', 2)
    add_para(doc, 'Module order.controller.js — hàm createOrder thực hiện toàn bộ quy trình đặt hàng trong một transaction, đảm bảo tính nhất quán dữ liệu.', indent=True)
    add_code_block(doc, 'Listing B.1 — Hàm createOrder (trích đoạn)',
                   read_source('backend/src/controllers/order.controller.js', 1, 110),
                   'backend/src/controllers/order.controller.js')

    add_heading(doc, 'Phụ lục C: Chatbot AI tích hợp Google Gemini', 2)
    add_para(doc, 'Module chat.controller.js — hàm sendMessage xây dựng system prompt từ catalog sản phẩm thực tế, gọi Gemini API và lưu lịch sử hội thoại.', indent=True)
    add_code_block(doc, 'Listing C.1 — Hàm sendMessage (trích đoạn)',
                   read_source('backend/src/controllers/chat.controller.js', 60, 153),
                   'backend/src/controllers/chat.controller.js')

    add_heading(doc, 'Phụ lục D: Frontend — Chat Widget (React/TypeScript)', 2)
    add_para(doc, 'Component ChatWidget.tsx quản lý phiên chat phía client, khởi tạo session và gửi tin nhắn qua Axios.', indent=True)
    add_code_block(doc, 'Listing D.1 — Khởi tạo phiên chat (trích đoạn)',
                   read_source('frontend/src/components/chat/ChatWidget.tsx', 1, 80),
                   'frontend/src/components/chat/ChatWidget.tsx')

    add_heading(doc, 'Phụ lục E: Cấu trúc cơ sở dữ liệu (Schema SQL)', 2)
    add_para(doc, 'File schema.sql định nghĩa cấu trúc các bảng cốt lõi của hệ thống.', indent=True)
    add_code_block(doc, 'Listing E.1 — Bảng users, products, orders (trích đoạn)',
                   read_source('database/schema.sql', 12, 134),
                   'database/schema.sql')

    add_heading(doc, 'Phụ lục F: Middleware xác thực JWT', 2)
    add_para(doc, 'Middleware authMiddleware kiểm tra Bearer token trên header Authorization, giải mã JWT và gắn thông tin user vào request.', indent=True)
    add_code_block(doc, 'Listing F.1 — authMiddleware',
                   read_source('backend/src/middlewares/auth.middleware.js', 1, 50),
                   'backend/src/middlewares/auth.middleware.js')

    add_heading(doc, 'Phụ lục G: Cấu hình API Frontend (Axios)', 2)
    add_para(doc, 'File api.ts cấu hình Axios instance, tự động gắn JWT token vào mọi request.', indent=True)
    add_code_block(doc, 'Listing G.1 — Axios interceptor',
                   read_source('frontend/src/lib/api.ts', 1, 40),
                   'frontend/src/lib/api.ts')

    add_para(doc, 'Ghi chú: Toàn bộ mã nguồn dự án bao gồm hơn 15.000 dòng code, được tổ chức theo cấu trúc monorepo với thư mục backend/ (Node.js/Express), frontend/ (Next.js/TypeScript) và database/ (MySQL schema). Sinh viên có thể nộp kèm USB/CD chứa mã nguồn đầy đủ hoặc link repository GitHub theo yêu cầu của Khoa.', indent=True)

    doc.save(OUTPUT)
    print(f'Saved: {OUTPUT}')

if __name__ == '__main__':
    build_document()
